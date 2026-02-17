import json
import hashlib
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document

from app.core.config import BASE_DIR
from app.models.schemas import DPTemplateInfo

CATALOG_DIR = BASE_DIR / "dp_catalog"
CATALOG_INDEX = CATALOG_DIR / "_index.json"


def _ensure_catalog():
    CATALOG_DIR.mkdir(parents=True, exist_ok=True)
    if not CATALOG_INDEX.exists():
        CATALOG_INDEX.write_text("[]", encoding="utf-8")


def _load_index() -> list[dict]:
    _ensure_catalog()
    return json.loads(CATALOG_INDEX.read_text(encoding="utf-8"))


def _save_index(index: list[dict]):
    _ensure_catalog()
    CATALOG_INDEX.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    import io
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def extract_text_from_docx(file_bytes: bytes) -> str:
    import io
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def save_dp_template(filename: str, file_bytes: bytes, name: str) -> DPTemplateInfo:
    _ensure_catalog()

    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        text_content = extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        text_content = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Formato nao suportado: {ext}. Use PDF ou DOCX.")

    template_id = hashlib.md5(file_bytes).hexdigest()[:12]

    original_path = CATALOG_DIR / f"{template_id}_{filename}"
    original_path.write_bytes(file_bytes)

    md_path = CATALOG_DIR / f"{template_id}.md"
    md_path.write_text(text_content, encoding="utf-8")

    preview = text_content[:300].replace("\n", " ").strip()
    if len(text_content) > 300:
        preview += "..."

    index = _load_index()
    entry = {
        "id": template_id,
        "name": name or Path(filename).stem,
        "filename": filename,
        "original_file": original_path.name,
        "md_file": md_path.name,
        "preview": preview,
    }
    index = [e for e in index if e["id"] != template_id]
    index.append(entry)
    _save_index(index)

    return DPTemplateInfo(
        id=template_id,
        name=entry["name"],
        filename=filename,
        preview=preview,
    )


def list_dp_templates() -> list[DPTemplateInfo]:
    index = _load_index()
    return [
        DPTemplateInfo(
            id=e["id"],
            name=e["name"],
            filename=e["filename"],
            preview=e["preview"],
        )
        for e in index
    ]


def get_dp_template_content(template_id: str) -> str:
    md_path = CATALOG_DIR / f"{template_id}.md"
    if md_path.exists():
        return md_path.read_text(encoding="utf-8")
    raise FileNotFoundError(f"Template {template_id} nao encontrado.")


def delete_dp_template(template_id: str) -> bool:
    index = _load_index()
    entry = next((e for e in index if e["id"] == template_id), None)
    if not entry:
        return False

    for key in ("original_file", "md_file"):
        fpath = CATALOG_DIR / entry[key]
        if fpath.exists():
            fpath.unlink()

    index = [e for e in index if e["id"] != template_id]
    _save_index(index)
    return True
